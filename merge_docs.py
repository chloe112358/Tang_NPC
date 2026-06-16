"""
合併腳本：
1. 把「期末報告_0605_已改好」中的第二大段(17-44)、第四大段(65-86)、6.2節(135-140)
   替換為「唐代報告_定稿_final」的對應內容
2. 在替換的第四大段中，調換圖3(Render環境變數)和圖4(Render Web Service)的位置
"""

import copy
import sys
import docx
from docx.oxml.ns import qn
from lxml import etree

def copy_paragraph(para):
    """Deep copy a paragraph element."""
    return copy.deepcopy(para._element)

def swap_image_paragraphs(tang_paragraphs, para_idx_a, para_idx_b):
    """
    Swap two pairs of paragraphs (image + caption) in the tang paragraphs list.
    para_idx_a: index of the image paragraph for figure A
    para_idx_b: index of the image paragraph for figure B
    We swap (para_idx_a, para_idx_a+1) with (para_idx_b, para_idx_b+1)
    """
    # Get elements at these indices
    img_a = tang_paragraphs[para_idx_a]._element
    cap_a = tang_paragraphs[para_idx_a + 1]._element
    img_b = tang_paragraphs[para_idx_b]._element
    cap_b = tang_paragraphs[para_idx_b + 1]._element

    # Deep copy all 4
    img_a_copy = copy.deepcopy(img_a)
    cap_a_copy = copy.deepcopy(cap_a)
    img_b_copy = copy.deepcopy(img_b)
    cap_b_copy = copy.deepcopy(cap_b)

    # Replace: put B's content into A's position, and A's content into B's position
    parent = img_a.getparent()
    
    # Find positions
    pos_img_a = list(parent).index(img_a)
    pos_cap_a = list(parent).index(cap_a)
    pos_img_b = list(parent).index(img_b)
    pos_cap_b = list(parent).index(cap_b)

    # Replace elements
    parent[pos_img_a] = img_b_copy
    parent[pos_cap_a] = cap_b_copy
    parent[pos_img_b] = img_a_copy
    parent[pos_cap_b] = cap_a_copy

    print(f"Swapped paragraphs at positions {para_idx_a},{para_idx_a+1} with {para_idx_b},{para_idx_b+1}")


def get_body_element(doc):
    """Get the body element of a document."""
    return doc.element.body


def get_section_elements(doc, start_idx, end_idx):
    """
    Extract paragraph elements from doc between start_idx and end_idx (inclusive).
    Returns list of XML elements (deep copies).
    """
    paras = doc.paragraphs
    elements = []
    for i in range(start_idx, end_idx + 1):
        elements.append(copy.deepcopy(paras[i]._element))
    return elements


def delete_paragraphs_in_range(doc, start_idx, end_idx):
    """
    Delete paragraphs from start_idx to end_idx (inclusive) from the document.
    """
    paras = doc.paragraphs
    elements_to_remove = []
    for i in range(start_idx, end_idx + 1):
        elements_to_remove.append(paras[i]._element)
    
    body = doc.element.body
    for elem in elements_to_remove:
        body.remove(elem)


def insert_elements_before(doc, before_idx, elements):
    """
    Insert XML elements before the paragraph at before_idx.
    """
    paras = doc.paragraphs
    body = doc.element.body
    
    # Get the reference element
    ref_elem = paras[before_idx]._element
    ref_pos = list(body).index(ref_elem)
    
    for i, elem in enumerate(elements):
        body.insert(ref_pos + i, elem)


def main():
    print("Loading documents...")
    tang = docx.Document('唐代報告_定稿_final.docx')
    final = docx.Document('期末報告_0605_已改好.docx')

    # Print summary of structure to verify
    print("\n=== 唐代報告 段落數:", len(tang.paragraphs))
    print("=== 期末報告 段落數:", len(final.paragraphs))

    # =============================================
    # Step 1: Extract sections from Tang document
    # =============================================
    
    # 唐代報告的「第二大段」= 段落 0-30
    # (二、系統架構與功能設計，含2.1到2.4)
    tang_section2_start = 0
    tang_section2_end = 30
    
    # 唐代報告的「第四大段」= 段落 31-69
    # (四、實作方法，含4.1到4.5)
    tang_section4_start = 31
    tang_section4_end = 69
    
    # 唐代報告的「6.2節」= 段落 70-73
    tang_section62_start = 70
    tang_section62_end = 73
    
    print(f"\n提取唐代報告第二大段 (段落 {tang_section2_start}-{tang_section2_end})...")
    tang_s2_elements = get_section_elements(tang, tang_section2_start, tang_section2_end)
    
    print(f"提取唐代報告第四大段 (段落 {tang_section4_start}-{tang_section4_end})...")
    tang_s4_elements = get_section_elements(tang, tang_section4_start, tang_section4_end)
    
    print(f"提取唐代報告6.2節 (段落 {tang_section62_start}-{tang_section62_end})...")
    tang_s62_elements = get_section_elements(tang, tang_section62_start, tang_section62_end)
    
    # =============================================
    # Step 2: Swap figures 3 and 4 in extracted section 4 elements
    # =============================================
    # In tang doc, para 58 = image for Render env vars, para 59 = its caption
    #              para 60 = image for Render Web Service, para 61 = its caption
    # These correspond to indices (58-31)=27 and (60-31)=29 in our extracted elements list
    
    # Let's verify by looking at captions
    print("\n=== 驗證要調換的圖片 ===")
    for idx in [27, 28, 29, 30]:
        para_text = tang.paragraphs[tang_section4_start + idx].text[:80]
        print(f"  提取元素[{idx}] (tang[{tang_section4_start+idx}]): {para_text}")
    
    # Swap figure 3 and figure 4 in the extracted elements
    # tang_s4_elements[27] = image para (tang[58])
    # tang_s4_elements[28] = caption para (tang[59]) "圖：Render 平台環境變數設定"
    # tang_s4_elements[29] = image para (tang[60])  
    # tang_s4_elements[30] = caption para (tang[61]) "圖：Render Web Service 設定介面"
    
    print("\n調換圖3和圖4 (Render環境變數 <-> Render Web Service)...")
    # Swap: element 27 with 29, and element 28 with 30
    tang_s4_elements[27], tang_s4_elements[29] = tang_s4_elements[29], tang_s4_elements[27]
    tang_s4_elements[28], tang_s4_elements[30] = tang_s4_elements[30], tang_s4_elements[28]
    print("調換完成！")
    
    # =============================================
    # Step 3: Replace sections in final document
    # Note: We need to do replacements from BOTTOM to TOP to preserve indices
    # =============================================
    
    # 期末報告的「6.2節」= 段落 135-140
    # 期末報告的「第四大段」= 段落 65-86
    # 期末報告的「第二大段」= 段落 17-44
    
    final_section62_start = 135
    final_section62_end = 140
    
    final_section4_start = 65
    final_section4_end = 86
    
    final_section2_start = 17
    final_section2_end = 44
    
    # --- Replace 6.2 section (bottom-most first) ---
    print(f"\n替換期末報告6.2節 (段落 {final_section62_start}-{final_section62_end})...")
    
    # Find the paragraph after the section to insert before
    # Para 141 is 6.3, which comes after 6.2
    # After deletion, we need to insert before what was at 141
    
    # Get the element after section 62 (para 141) before deletion
    after_s62_elem = final.paragraphs[final_section62_end + 1]._element
    
    # Delete 6.2 section
    delete_paragraphs_in_range(final, final_section62_start, final_section62_end)
    
    # Insert Tang's 6.2 before the element that was at para 141
    body = final.element.body
    ref_pos = list(body).index(after_s62_elem)
    for i, elem in enumerate(tang_s62_elements):
        body.insert(ref_pos + i, copy.deepcopy(elem))
    print(f"  插入了 {len(tang_s62_elements)} 個段落")
    
    # Re-count paragraphs after modification
    print(f"  操作後期末報告段落數: {len(final.paragraphs)}")
    
    # --- Replace section 4 ---
    # Need to find new indices after 6.2 replacement
    # Find section 4 by scanning for "四、實作方法" heading
    new_s4_start = None
    new_s4_end = None
    for i, para in enumerate(final.paragraphs):
        t = para.text.strip()
        if para.style.name == 'Heading 1' and '實作方' in t:
            new_s4_start = i
        if new_s4_start and para.style.name == 'Heading 1' and '使用者' in t:
            new_s4_end = i - 1
            # Go back to find last non-empty
            while not final.paragraphs[new_s4_end].text.strip() and new_s4_end > new_s4_start:
                new_s4_end -= 1
            break
    
    print(f"\n找到期末報告第四大段: 段落 {new_s4_start}-{new_s4_end}")
    
    # Get element after section 4
    after_s4_elem = final.paragraphs[new_s4_end + 1]._element
    
    # Delete section 4
    delete_paragraphs_in_range(final, new_s4_start, new_s4_end)
    
    # Insert Tang's section 4 (with swapped figures)
    ref_pos = list(body).index(after_s4_elem)
    for i, elem in enumerate(tang_s4_elements):
        body.insert(ref_pos + i, copy.deepcopy(elem))
    print(f"  插入了 {len(tang_s4_elements)} 個段落")
    print(f"  操作後期末報告段落數: {len(final.paragraphs)}")
    
    # --- Replace section 2 ---
    new_s2_start = None
    new_s2_end = None
    for i, para in enumerate(final.paragraphs):
        t = para.text.strip()
        if para.style.name == 'Heading 1' and '系統架構' in t:
            new_s2_start = i
        if new_s2_start and para.style.name == 'Heading 1' and '三、' in t:
            new_s2_end = i - 1
            while not final.paragraphs[new_s2_end].text.strip() and new_s2_end > new_s2_start:
                new_s2_end -= 1
            break
    
    print(f"\n找到期末報告第二大段: 段落 {new_s2_start}-{new_s2_end}")
    
    # Get element after section 2
    after_s2_elem = final.paragraphs[new_s2_end + 1]._element
    
    # Delete section 2
    delete_paragraphs_in_range(final, new_s2_start, new_s2_end)
    
    # Insert Tang's section 2
    ref_pos = list(body).index(after_s2_elem)
    for i, elem in enumerate(tang_s2_elements):
        body.insert(ref_pos + i, copy.deepcopy(elem))
    print(f"  插入了 {len(tang_s2_elements)} 個段落")
    print(f"  操作後期末報告段落數: {len(final.paragraphs)}")
    
    # =============================================
    # Step 4: Save
    # =============================================
    output_path = '期末報告_0605_已改好_合併版.docx'
    final.save(output_path)
    print(f"\n[OK] 已儲存合併後的文件：{output_path}")
    
    # Verify structure
    print("\n=== 驗證最終文件結構 ===")
    for i, para in enumerate(final.paragraphs):
        if para.style.name.startswith('Heading'):
            print(f"[{i}] {para.style.name}: {para.text[:80]}")


if __name__ == '__main__':
    main()
